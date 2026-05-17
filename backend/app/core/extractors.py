"""Knowledge document text extractors.

Single point of truth for converting an uploaded artifact (bytes + filename
+ optional MIME) into the plain text that downstream chunking, embedding,
and graph extraction expect. Adding a new file format means adding one
function and one entry to ``_DISPATCH``.

Design notes
------------
- We dispatch primarily on **filename suffix** because that's both the
  most stable client-supplied hint and the only information we have when
  re-running the pipeline from a stored MinIO blob during retry.
- ``mime`` is consulted only as a tiebreaker when the suffix is unknown
  (URL ingestion of a server that returns a useful Content-Type but a
  meaningless URL path).
- Errors raise ``UnsupportedSourceError`` (415-shaped) or
  ``ExtractionError`` (422-shaped) so the API layer can map them to the
  correct HTTP status without sniffing exception messages.
- Each extractor is **synchronous and CPU-bound**; callers run them in a
  thread executor exactly like the existing ``extract_pdf_text`` does.
"""
from __future__ import annotations

import io
import os
from dataclasses import dataclass
from typing import Callable


class UnsupportedSourceError(ValueError):
    """Raised when the file extension/mime is not in the supported set."""


class ExtractionError(ValueError):
    """Raised when a supported source fails to yield usable text."""


# Suffix-only canonical type tag stored on KnowledgeDoc.source_type.
# Keep in sync with the spec at:
#   docs/superpowers/specs/2026-05-09-knowledge-multi-source-design.md
SUPPORTED_SUFFIXES: tuple[str, ...] = (
    ".pdf",
    ".txt",
    ".md",
    ".markdown",
    ".html",
    ".htm",
    ".docx",
)


# ---------------------------------------------------------------------------
# Individual extractors
# ---------------------------------------------------------------------------


def _extract_pdf(content: bytes) -> str:
    # Local import keeps the module importable in tests where pymupdf is
    # not installed (and matches the lazy-import pattern in knowledge.py).
    import fitz  # type: ignore

    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"PyMuPDF failed to open PDF: {exc}") from exc

    text_parts: list[str] = []
    # Same 200-page cap as the legacy path: prevents OOM on 500-page books
    # while still covering all real product manuals we've seen.
    MAX_PAGES = 200
    try:
        for i, page in enumerate(doc):
            if i >= MAX_PAGES:
                break
            text_parts.append(page.get_text())
    finally:
        doc.close()
    return "".join(text_parts)


def extract_pdf_page_images(
    content: bytes,
    max_pages: int = 50,
    dpi: int = 150,
) -> list[dict]:
    """Extract page images from a PDF for multimodal embedding.

    Returns a list of dicts, each with:
        - ``page``: 0-based page index
        - ``image_base64``: base64-encoded PNG data URI
        - ``text``: any extractable text on the page (may be empty)

    Useful for scanned/image-only PDFs where ``_extract_pdf`` returns empty text.
    Uses PyMuPDF to render each page to a PNG image.
    """
    import fitz  # type: ignore

    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise ExtractionError(f"PyMuPDF failed to open PDF: {exc}") from exc

    pages: list[dict] = []
    try:
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            # Render page to pixmap at requested DPI
            pix = page.get_pixmap(dpi=dpi)
            png_bytes = pix.tobytes("png")
            b64 = _encode_image_base64(png_bytes)
            text = page.get_text() or ""
            pages.append({
                "page": i,
                "image_base64": f"data:image/png;base64,{b64}",
                "text": text.strip(),
            })
    finally:
        doc.close()

    return pages


def _encode_image_base64(image_bytes: bytes) -> str:
    import base64
    return base64.b64encode(image_bytes).decode("ascii")


def _decode_text_bytes(content: bytes) -> str:
    """Decode arbitrary bytes to str.

    Strategy:
    1. Strip UTF-8 BOM if present.
    2. Try strict UTF-8 — succeeds for the overwhelming majority of
       modern documents.
    3. For larger non-UTF8 buffers (>=64 bytes) consult
       ``charset-normalizer``; it has enough material to be reliable.
    4. Fall back to GB18030 — a superset of GBK / GB2312 that covers
       essentially every legacy Chinese product manual we'll see in
       practice. This catches the short-input case where
       charset-normalizer is statistically unreliable.
    5. Last resort: lossy UTF-8 replace so we never crash the pipeline.
    """
    if content.startswith(b"\xef\xbb\xbf"):
        return content[3:].decode("utf-8", errors="replace")

    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        pass

    if len(content) >= 64:
        try:
            from charset_normalizer import from_bytes
            best = from_bytes(content).best()
            if best is not None:
                return str(best)
        except ImportError:
            pass

    # GB18030 reliably decodes pure ASCII, GBK, GB2312, and most CJK
    # byte streams. It tolerates malformed sequences less than UTF-8
    # but doesn't error on anything realistic for Chinese docs.
    try:
        return content.decode("gb18030")
    except UnicodeDecodeError:
        return content.decode("utf-8", errors="replace")


def _extract_txt(content: bytes) -> str:
    return _decode_text_bytes(content)


def _extract_markdown(content: bytes) -> str:
    # Treat markdown as plain text. Stripping syntax (#, **, etc.) would
    # actually hurt the embedder since the symbols carry weak semantic
    # signal and headers boost section relevance.
    return _decode_text_bytes(content)


# Tags that almost never carry useful body text — strip them outright so
# noise like menus and ads don't dilute embeddings.
_HTML_NOISE_TAGS = ("script", "style", "noscript", "nav", "header", "footer", "aside", "form")


def _extract_html(content: bytes) -> str:
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError as exc:
        raise ExtractionError("beautifulsoup4 not installed") from exc

    text = _decode_text_bytes(content)
    # `lxml` parser is much faster on large pages and tolerates malformed
    # HTML better than the stdlib parser. Fall back if lxml is missing
    # (matters only in lean test environments).
    try:
        soup = BeautifulSoup(text, "lxml")
    except Exception:  # noqa: BLE001
        soup = BeautifulSoup(text, "html.parser")

    for tag_name in _HTML_NOISE_TAGS:
        for tag in soup.find_all(tag_name):
            tag.decompose()

    body = soup.body or soup
    # Use \n as separator so paragraph chunks survive in the chunker that
    # splits on blank lines.
    return body.get_text(separator="\n", strip=True)


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise ExtractionError("python-docx not installed") from exc

    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"python-docx failed to open document: {exc}") from exc

    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables carry essential spec-sheet data in product docs; flatten cell
    # by cell separated with " | " so the embedder still sees them as
    # rows.
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class _Spec:
    source_type: str  # canonical tag persisted in DB
    extractor: Callable[[bytes], str]


_DISPATCH: dict[str, _Spec] = {
    ".pdf": _Spec("pdf", _extract_pdf),
    ".txt": _Spec("txt", _extract_txt),
    ".md": _Spec("md", _extract_markdown),
    ".markdown": _Spec("md", _extract_markdown),
    ".html": _Spec("html", _extract_html),
    ".htm": _Spec("html", _extract_html),
    ".docx": _Spec("docx", _extract_docx),
}


# Optional MIME → suffix fallback for URL ingestion where the URL has no
# meaningful path (e.g. ``https://api.example.com/doc?id=42``).
_MIME_FALLBACK: dict[str, str] = {
    "application/pdf": ".pdf",
    "text/plain": ".txt",
    "text/markdown": ".md",
    "text/html": ".html",
    "application/xhtml+xml": ".html",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
}


def normalize_suffix(filename: str) -> str:
    """Lower-cased extension including the dot, or ``''`` if none."""
    return os.path.splitext(filename or "")[1].lower()


def detect_source_type(filename: str, mime: str | None = None) -> str:
    """Return the canonical ``source_type`` tag without doing extraction.

    Raises ``UnsupportedSourceError`` if neither the suffix nor the MIME
    are recognized. Cheap to call from request handlers for early 415s.
    """
    suffix = normalize_suffix(filename)
    if suffix in _DISPATCH:
        return _DISPATCH[suffix].source_type

    if mime:
        mime_main = mime.split(";", 1)[0].strip().lower()
        suffix_guess = _MIME_FALLBACK.get(mime_main)
        if suffix_guess and suffix_guess in _DISPATCH:
            return _DISPATCH[suffix_guess].source_type

    raise UnsupportedSourceError(
        f"Unsupported document type: filename={filename!r} mime={mime!r}. "
        f"Supported: {', '.join(SUPPORTED_SUFFIXES)}"
    )


def extract_text(content: bytes, *, filename: str, mime: str | None = None) -> str:
    """Bytes → text. Dispatches on suffix first, then MIME.

    Raises ``UnsupportedSourceError`` for unknown types and
    ``ExtractionError`` when extraction itself fails or yields no text.
    """
    suffix = normalize_suffix(filename)
    spec: _Spec | None = _DISPATCH.get(suffix)

    if spec is None and mime:
        mime_main = mime.split(";", 1)[0].strip().lower()
        suffix_guess = _MIME_FALLBACK.get(mime_main)
        if suffix_guess:
            spec = _DISPATCH.get(suffix_guess)

    if spec is None:
        raise UnsupportedSourceError(
            f"Unsupported document type: filename={filename!r} mime={mime!r}. "
            f"Supported: {', '.join(SUPPORTED_SUFFIXES)}"
        )

    text = spec.extractor(content)
    if not text or not text.strip():
        raise ExtractionError(
            "Extracted text is empty. The document may be image-only "
            "(scanned / OCR required), encrypted, or malformed."
        )
    return text


def derive_filename_from_mime(mime: str | None, fallback: str = "document") -> str:
    """Build a sensible filename for a URL fetch when the URL path is
    uninformative. We map MIME → suffix using ``_MIME_FALLBACK`` so the
    rest of the pipeline keeps a real extension to dispatch on.
    """
    if mime:
        mime_main = mime.split(";", 1)[0].strip().lower()
        suffix = _MIME_FALLBACK.get(mime_main)
        if suffix:
            return f"{fallback}{suffix}"
    return f"{fallback}.html"  # safest default for arbitrary web URL
